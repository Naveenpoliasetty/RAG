from docx import Document
import requests
import os
from typing import List, Dict, Optional
from pydantic import BaseModel


class Experience(BaseModel):
    client_name: str
    duration: str
    job_role: str
    responsibilities: List[str]
    environment: Optional[str] = None


class Resume(BaseModel):
    name: str
    designation: str
    phone_number: Optional[str] = None
    email: Optional[str] = None
    url: Optional[List[str]] = None
    professional_summary: List[str]
    technical_skills: Dict[str, List[str]]
    experiences: List[Experience]
    education: List[str]


def parse_resume(docx_path: str) -> Resume:
    doc = Document(docx_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    idx = 0

    # -------------------------
    # HEADER
    # -------------------------

    name = paragraphs[idx]
    idx += 1

    designation = paragraphs[idx]
    idx += 1

    # -------------------------
    # CONTACT
    # -------------------------

    phone_number = None
    email = None
    urls = None

    contact_line = paragraphs[idx]
    idx += 1

    parts = contact_line.split("\t")

    for part in parts:
        if part.startswith("Phone:"):
            phone_number = part.replace("Phone:", "").strip()
        elif part.startswith("Email:"):
            email = part.replace("Email:", "").strip()
        elif part.startswith("URL:"):
            urls = [u.strip() for u in part.replace("URL:", "").split(",")]

    # -------------------------
    # HELPER
    # -------------------------

    def move_to_section(title):
        nonlocal idx
        while idx < len(paragraphs) and paragraphs[idx].upper() != title.upper():
            idx += 1
        idx += 1

    # -------------------------
    # PROFESSIONAL SUMMARY
    # -------------------------

    move_to_section("PROFESSIONAL SUMMARY")

    professional_summary = []

    while idx < len(paragraphs) and not paragraphs[idx].isupper():
        professional_summary.append(paragraphs[idx])
        idx += 1

    # -------------------------
    # TECHNICAL SKILLS
    # -------------------------


    # safer table handling
    technical_skills = {}

    if doc.tables:
        skills_table = doc.tables[0]

        for row in skills_table.rows:
            category = row.cells[0].text.replace(":", "").strip()
            values = [v.strip() for v in row.cells[1].text.split(",")]
            technical_skills[category] = values

    # -------------------------
    # EXPERIENCE
    # -------------------------

    move_to_section("PROFESSIONAL EXPERIENCE")

    experiences = []

    while idx < len(paragraphs):

        if paragraphs[idx].upper() == "EDUCATION":
            break

        header = paragraphs[idx]
        idx += 1

        if "\t" not in header:
            continue

        client_name, duration = header.split("\t")

        role_line = paragraphs[idx]
        idx += 1

        job_role = role_line.replace("Role:", "").strip()

        responsibilities = []

        while (
            idx < len(paragraphs)
            and not paragraphs[idx].startswith("Environment:")
            and not paragraphs[idx].isupper()
        ):
            responsibilities.append(paragraphs[idx])
            idx += 1

        environment = None

        if idx < len(paragraphs) and paragraphs[idx].startswith("Environment:"):
            environment = paragraphs[idx].replace("Environment:", "").strip()
            idx += 1

        experiences.append(
            Experience(
                client_name=client_name.strip(),
                duration=duration.strip(),
                job_role=job_role,
                responsibilities=responsibilities,
                environment=environment
            )
        )

    # -------------------------
    # EDUCATION
    # -------------------------

    move_to_section("EDUCATION")

    education = []

    while idx < len(paragraphs):
        education.append(paragraphs[idx])
        idx += 1

    # -------------------------
    # RETURN PYDANTIC OBJECT
    # -------------------------

    return Resume(
        name=name,
        designation=designation,
        phone_number=phone_number,
        email=email,
        url=urls,
        professional_summary=professional_summary,
        technical_skills=technical_skills,
        experiences=experiences,
        education=education
    )


def download_resume_from_gcs(gcs_url: str, save_dir: str = ".") -> str:
    """
    Downloads a resume from GCS public URL and saves locally.

    Returns local file path.
    """

    response = requests.get(gcs_url, stream=True)
    response.raise_for_status()

    # Extract filename from URL
    filename = gcs_url.split("/")[-1]
    local_path = os.path.join(save_dir, filename)

    with open(local_path, "wb") as f:
        for chunk in response.iter_content(chunk_size=8192):
            if chunk:
                f.write(chunk)

    return local_path