from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_resume():
    doc = Document()
    
    # --- GLOBAL STYLES ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # ------- REMOVE TABLE BORDERS ------

    def remove_table_borders(table):
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for b in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    edge = OxmlElement(f'w:{b}')
                    edge.set(qn('w:val'), 'nil')
                    tcBorders.append(edge)
                tcPr.append(tcBorders)

    # ------- HELPER FUNCTIONS -------

    def add_bottom_border(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        add_bottom_border(p)

    def add_bullet_point(text):
        if text:
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    def add_experience_block(company, location, role, dates, bullets, environment=None):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        remove_table_borders(table)
        
        cell_1 = table.rows[0].cells[0]
        cell_1.width = Inches(5.0)
        p1 = cell_1.paragraphs[0]
        r1 = p1.add_run(company)
        r1.bold = True
        r1.font.size = Pt(11)
        if location:
            p1.add_run(f", {location}").bold = True

        cell_2 = table.rows[0].cells[1]
        cell_2.width = Inches(2.5)
        p2 = cell_2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if dates:
            p2.add_run(dates).bold = True

        if role:
            p_role = doc.add_paragraph()
            r_role = p_role.add_run(f"Role: {role}")
            r_role.italic = True
        
        if bullets:
            for bullet in bullets:
                add_bullet_point(bullet)

        if environment:
            p_env = doc.add_paragraph()
            r_lbl = p_env.add_run("Environment: ")
            r_lbl.bold = True
            p_env.add_run(environment)

    # ------- HEADER -------

    h_p = doc.add_paragraph()
    h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_p.paragraph_format.space_before = Pt(0)   # move up
    h_p.paragraph_format.space_after = Pt(2)

    r_name = h_p.add_run("Aasreetha Saggu\n")
    r_name.bold = True
    r_name.font.size = Pt(18)

    r_title = h_p.add_run("SR. SAP Consultant\n")
    r_title.bold = True
    r_title.font.size = Pt(12)

    # Contact Table
    contact_table = doc.add_table(rows=1, cols=3)
    contact_table.autofit = False
    remove_table_borders(contact_table)

    cell_l = contact_table.rows[0].cells[0]
    cell_l.width = Inches(2.5)
    cell_l.paragraphs[0].add_run("Email: aasreetharao@gmail.com").bold = True

    contact_table.rows[0].cells[1].width = Inches(3.0)

    cell_r = contact_table.rows[0].cells[2]
    cell_r.width = Inches(2.5)
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.add_run("Phone: 469-901-5194").bold = True

    # ------- PROFESSIONAL SUMMARY -------

    add_section_header("Professional Summary")

    summary_points = [
        "Over 10+ years of specialized expertise in SAP Production Planning (PP) and Production Execution (PE).",
        "Extensive experience in designing, configuring, and supporting SAP ERP systems.",
        "Strong focus on Manufacturing Execution Systems (MES) and Quality Management (QM) modules.",
        "Expertise in SAP S/4HANA, PP/PE, MM, SD, WM, and QM integration.",
        "Skilled in MRP, Capacity Planning, Shop Floor Control, and Batch Management."
    ]

    for point in summary_points:
        add_bullet_point(point)

    # 3. TECHNICAL SKILLS
    # [cite_start]Data derived from source [cite: 27]
    skills_data = [
        ("SAP Modules:", "Sales and Distribution (SD), SAP QM (Quality Management), Materials Management (MM), Logistics Execution (LE), SAP PP/PE, SAP WM, SAP S/4HANA, SAP LE, SAP PMR"),
        ("Configuration:", "Production Orders, Work Centers, Routings, Bill of Materials (BOM), Make-to-Stock (MTS), Make-to-Order (MTO), Order-to-Cash (OTC), ATP Checks, EDI Transactions: 850, 855, 856, 940, and 810"),
        ("Integration:", "Cross-functional integration of SD with MM modules, Integration with Logistics Execution (LE) and Warehouse Management (WM), SAP PP"),
        ("Project Management:", "Project Planning, Testing, Production Support, Change Management, Material Master, Vendor Master, Purchase Info Records, Bills of Materials (BOM), Work Centers."),
        ("Interface Technologies:", "EDI, IDocs, ALE, RFCs"),
        ("Production Planning:", "Routings, Production Versions, Cost Collectors, Production Orders, MRP Run Configuration"),
        ("Tools:", "SAP Solution Manager, SAP PowerDesigner, MS Excel, MS Visio, MS Word, MS PowerPoint, SAP Fiori, UI5"),
        ("Documentation:", "Functional Requirement Specifications (FRSs), Dataflow Diagrams, Functional Specifications, Unit Test Scripts"),
        ("Interfacing with Third-party Systems:", "Stay in Front, Salesforce, Quofor")
    ]

    if skills_data:
        add_section_header("Technical Skills")
        skill_table = doc.add_table(rows=0, cols=2)
        skill_table.style = 'Table Grid' # Using grid for clear separation as requested by "Table" in source
        
        for category, items in skills_data:
            if items: # Only add row if data exists
                row = skill_table.add_row()
                # Category Column
                c1 = row.cells[0]
                c1.width = Inches(2.0)
                r_cat = c1.paragraphs[0].add_run(category)
                r_cat.bold = True
                r_cat.font.size = Pt(10)
                
                # Details Column
                c2 = row.cells[1]
                c2.width = Inches(5.5)
                c2.paragraphs[0].add_run(items).font.size = Pt(10)

    # 4. PROFESSIONAL EXPERIENCE
    add_section_header("Professional Experience")

    # [cite_start]Experience Data List (Extracted from source [cite: 28-142])
    # This structure allows you to easily handle missing data by leaving fields as None or ""
    experience_list = [
        {
            "client": "Best buy",
            "loc": "Richfield, MN",
            "dates": "Jul 2021 – Present",
            "role": "Sr. SAP Techno Functional Consultant",
            "env": "SAP QM/PP, MIGO, ABAP, SAP Query, IDOC, EDI, SAP GUI, SAP QM tools, STO, logistics invoice verification, and pricing procedures management, TM, LE, EWM, RICEFW, SAP SD, MM, WM, SAP PMR, Master Data, Order to Cash (OTC), EDI, IDocs, Workflow, SAP Landscape ABAP, Change Management, UDF, SAP Framework",
            "tasks": [
                "Configured Production Execution (PE) processes for Make-to-Order and Make-to-Stock manufacturing, enabling seamless production workflows.",
                "Integrated SAP PP/PE with Warehouse Management (WM) and Quality Management (QM) for synchronized inventory tracking and quality inspections.",
                "Configured and optimized MRP types (MRP, Consumption-Based Planning, Reorder Point Planning) for diverse manufacturing scenarios.",
                "Enhanced MRP functionalities by configuring planning parameters like lot sizing, safety stock, and planning horizons.",
                "Configured and managed production versions, BOMs, and routings to support diverse manufacturing needs.",
                "Implemented Transfer Orders/Requests (TR/TO) for optimized warehouse movements and seamless inventory flow.",
                "Developed and executed custom BAPIs, RFCs, and user exits to meet specific PP/PE functional needs.",
                "Linkage between inspection lots and production orders for efficient tracking and management of quality inspections.",
                "Configured end-to-end OTC processes, including setup for sales order types, billing, and returns management.",
                "Performed single - item, single-item/multi-level MRP runs and monitored MRP/Stock requirement list-individual and collective display.",
                "Configured Production Scheduling Profiles, Order Confirmation Strategies, and integration points for Advanced ATP (Available-to-Promise) checks.",
                "Configured Purchasing (MM) processes, including purchase requisitions, requests for quotations (RFQs), and purchase orders.",
                "Customized release strategies and approval workflows for procurement cycles using SAP Workflow and BRF+ rules.",
                "Integrated MRP-driven procurement planning with purchasing workflows.",
                "Created and managed vendor evaluation reports, assessing supplier performance based on delivery timelines, quality, and costs.",
                "Designed and optimized material planning strategies, supporting Make-to-Order (MTO) and Make-to-Stock (MTS).",
                "Provided hands-on configuration for purchase scheduling agreements and contract management.",
                "Collaborated with cross-functional teams, including finance, warehouse, and logistics.",
                "Performed data migration and validation for procurement master data using LSMW.",
                "Developed custom reports and dashboards for purchasing trends, order history, and supplier performance using SAP Analytics Cloud and SAP BW.",
                "Conducted training workshops and created training manuals for end-users.",
                "Setup of sampling procedures in SAP QM that define how sample quantities are selected during the production process.",
                "Provided ongoing support and optimization for production planning (PP) in SAP S/4HANA.",
                "Utilized ALE/IDocs to streamline communications between SAP and third-party systems.",
                "Configuring MRP groups for import, local consignment and finished goods as per business requirement.",
                "Customized SAP PP features for make-to-order and make-to-stock production.",
                "Expert in customizing SAP PP/QM modules, including Production Order Types, Work Center Configuration, Quality Inspection Plans, and Control Keys.",
                "Implemented Problem Management procedures, tracking root causes and implementing corrective actions.",
                "Created MRP element through MRP runs, basic and lead time scheduling.",
                "Also evaluating MRP results via MRP lists and Stock requirement list.",
                "Implemented process improvements and enhancements within SAP PM.",
                "Worked on Procure to Pay (P2P) including manual and automatic creation of Purchase orders, invoice verification.",
                "Performed unit testing, UAT on various process flows for both PI and MM modules.",
                "Implemented ATP (Available to Promise) functionality within SCM."
            ]
        },
        {
            "client": "Caterpillar Inc",
            "loc": "Irving, Texas",
            "dates": "May 2019 – Jun 2021",
            "role": "SAP Techno Functional Consultant",
            "env": "SAP Test Acceleration and Optimization (TAO), SAP ERP, MM, PM, SD, PP, WM, QM, SAP Solution Manager, Microsoft Excel, JIRA, Service Now, TM, LE, EWM, SAP PMR, Master Data, Order to Cash (OTC), EDI, IDocs, Workflow, SAP Landscape ABAP, Change Management, UDF, Problem Management, LSMW, BDC, SAP GUI, BOM, MRP, IDOC, EDI.",
            "tasks": [
                "Designed and configured SAP S/4HANA Production Planning and Execution (PP/PE), aligning with Caterpillar’s manufacturing workflows.",
                "Implemented Material Requirements Planning (MRP) to optimize raw material utilization and streamline production schedules.",
                "Configured SAP PP (Production Planning) for discrete and repetitive manufacturing processes.",
                "Integrated SAP PP/PE with Manufacturing Execution Systems (MES), enabling real-time shop floor data updates.",
                "Customized BOM (Bill of Materials) for multi-level production requirements.",
                "Configured SAP S/4HANA for discrete and repetitive manufacturing processes.",
                "Configured and maintained EDI interfaces for transactions like 850, 855, 846, 856, 940, 856SC, and 810.",
                "Designed custom IDocs for 856 ASN and 810 transactions.",
                "Enabled batch management and integrated SAP QM processes for ensuring high-quality standards.",
                "Set up master data structures, including Material Master, BOMs, Work Centers, and Routings.",
                "Conducted detailed gap analysis for OTC processes, configured customer master data, pricing, and ATP settings.",
                "Integrated SAP PP with Warehouse Management (WM) and SAP QM.",
                "Configuration of Make-to-Stock (MTS) and Make-to-Order (MTO) strategies.",
                "Implemented Material Requirements Planning (MRP) in SAP S/4HANA.",
                "Experience with SAP User Exits, BADIs, and ABAP Debugging to customize and enhance SAP PP/QM functionality.",
                "Strong in maintaining, cleansing, and managing Master Data.",
                "Conducted requirement gathering sessions with stakeholders to analyze procurement processes.",
                "Provided UAT testing, validation, and cutover support for procurement-related workflows.",
                "Integrated Purchasing (MM) with Inventory Management (IM) and Warehouse Management (WM).",
                "Delivered training documents and hands-on end-user training.",
                "Supported Supplier Evaluation Processes, ensuring compliance with performance metrics.",
                "Automated Purchase Order Approvals and Release Strategies using BRF+ rules.",
                "Utilized UDF to enhance demand forecasting accuracy.",
                "Utilized ALE and IDocs to facilitate real-time data transfer.",
                "Implemented Lead-to-Cash processes in SAP Sales Cloud."
            ]
        },
        {
            "client": "Fortis Healthcare",
            "loc": "Gurgaon, India",
            "dates": "Aug 2017 – Jan 2019",
            "role": "SAP Techno Functional Consultant",
            "env": "SAP GUI, SAP MM, SAP PP, SAP Solution Manager, SAP ERP, SAP QM, SAP PP, Microsoft Excel, TM, LE, EWM, Microsoft Project, JIRA, Git, Slack, Microsoft Word, ITSM, Camstar, SAP PMR, Master Data, Order to Cash (OTC), EDI, IDocs, Workflow, SAP Landscape ABAP, SAP Query, Microsoft Excel, SAP Testing Tools, SAP GUI, BOM, MRP",
            "tasks": [
                "Led the implementation of SAP S/4HANA PP/PE modules to support hospital equipment manufacturing.",
                "Configured and maintained batch tracking, inspection lots, and quality notifications.",
                "Designed and implemented MRP processes to streamline high-volume procurement of pharmaceuticals.",
                "Managed Material Requirements Planning (MRP) by handling processing keys such as NETCH, NETPL, and NEUPL.",
                "Collected and documented business requirements, preparing comprehensive Business Requirement Documents (BRD).",
                "Defined and implemented release procedures for purchase requisitions and external purchasing documents.",
                "Conducted user acceptance testing (UAT) on Electronic Data Interchange (EDI) transaction sets.",
                "Configured warehouse management (WM) settings to optimize medical inventory handling.",
                "Developed functional specifications for seamless integration between SAP and Camstar.",
                "Configured BAPIs and RFCs to facilitate real-time communication between SAP and third-party logistics providers (3PL).",
                "Enhanced Production Planning and Detailed Scheduling (PP/DS) settings.",
                "Configured SAP FI settings for integration with Credit Management.",
                "Integrated SAP PP and SAP QM with Material Management (MM).",
                "Configured transfer orders and transfer requirements in SAP WM.",
                "Configured Schedule Margin Keys and strategy groups to optimize Make-to-Order (MTO).",
                "Configured customer master data and partner determination settings within SAP SD.",
                "Managed IDOC and EDI processes for healthcare logistics.",
                "Configured MRP parameters for lot sizing and safety stock."
            ]
        },
        {
            "client": "Amway Corp",
            "loc": "New Delhi, India",
            "dates": "Jun 2014– Jul 2017",
            "role": "SAP Analyst",
            "env": "SAP Sales and Distribution (SD), SAP GUI, SAP Business Process Documentation, MS Word, MS PowerPoint, MS Visio, ABAP, RICEFW, Unit testing, Project management.",
            "tasks": [
                "Worked on SAP Sales and Distribution project. Studied the existing business processes with users.",
                "Mapped existing business processes to SAP business processes.",
                "Collaborated with stakeholders to gather user requirements and prepare business blueprints for SAP MM module.",
                "Actively contributed as a team member in the successful implementation of the SAP MM module.",
                "Responsible for several project phases including collecting business requirements and design specifications.",
                "Worked in the SD processes including creating sales orders, quotations, and master data.",
                "Created ABAP query reports to meet client’s reporting requirements.",
                "Supported end-users testing during the implementation.",
                "Initiated the new user training in numerous customized transactions.",
                "Configured the system to work for BOM processing.",
                "Worked on configuration required for EDI-810/850/856 (ASN) including partner profile, message controls.",
                "Configuration of pricing agreements including promotions and sales deals.",
                "Preparing business process documents for the sales and distribution activities.",
                "Prepared comprehensive business process documents using MS Word and MS PowerPoint.",
                "Prepared functional specification for the development of Reports, Interface, Conversion, Enhancement, and Forms.",
                "Prepared unit test scripts and checklist to be included in the functional specification.",
                "Configured customer master and partner determination."
            ]
        }
    ]

    for job in experience_list:
        add_experience_block(
            job.get("client"),
            job.get("loc"),
            job.get("role"),
            job.get("dates"),
            job.get("tasks"),
            job.get("env")
        )

    # 5. EDUCATION
    # [cite_start]Data from source [cite: 144, 145]
    edu_school = "Jawaharlal Nehru Technology University, Hyderabad, TS, India"
    edu_degree = "BTech in Computer Science and Engineering"
    edu_dates = "June 2010 - May 2014"

    if edu_school:
        add_section_header("Education")
        
        # Table for layout
        edu_table = doc.add_table(rows=0, cols=2)
        edu_table.autofit = False
        
        row = edu_table.add_row()
        # School
        c1 = row.cells[0]
        c1.width = Inches(5.0)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(edu_school)
        r1.bold = True
        
        # Degree (Line below school)
        if edu_degree:
            r1_deg = p1.add_run(f"\n{edu_degree}")
            r1_deg.italic = True

        # Date
        c2 = row.cells[1]
        c2.width = Inches(2.5)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if edu_dates:
            p2.add_run(edu_dates).bold = True

    # SAVE
    file_name = 'Aasreetha_Saggu_Resume_Refined.docx'
    doc.save(file_name)
    print(f"Document saved as {file_name}")

if __name__ == "__main__":
    create_resume()