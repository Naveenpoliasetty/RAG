from google.cloud import storage
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


# ------------------------------------------------------------
# 1. CREATE RESUME (.docx)
# ------------------------------------------------------------

def create_resume(resume_data: dict) -> str:
    doc = Document()

    # ---------- SHRINK ALL MARGINS ----------
    section = doc.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

    # Base style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    def add_bottom_border(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        add_bottom_border(p)

    def add_bullet_point(text):
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    # --------------------------------------------------------
    # HEADER — NAME + DESIGNATION
    # --------------------------------------------------------

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(resume_data["name"])
    r.bold = True
    r.font.size = Pt(18)

    desig_p = doc.add_paragraph()
    desig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d = desig_p.add_run(resume_data["designation"])
    d.italic = True
    d.font.size = Pt(11)

    # --------------------------------------------------------
    # CONTACT INFO (NO TABLES)
    # --------------------------------------------------------

    phone = resume_data.get("phone_number") or ""
    email = resume_data.get("email") or ""
    url_list = resume_data.get("url")
    url = ", ".join(url_list) if isinstance(url_list, list) else (url_list or "")

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tab stops for spacing
    ts = contact_p.paragraph_format.tab_stops
    ts.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.CENTER)
    ts.add_tab_stop(Inches(5), WD_TAB_ALIGNMENT.CENTER)

    contact_parts = []

    if phone:
        contact_parts.append(f"Phone: {phone}")

    if email:
        contact_parts.append(f"Email: {email}")

    if url:  # url can be None, empty string, or empty list
        contact_parts.append(f"URL: {url}")

    contact_text = "\t".join(contact_parts)

    if contact_text:  # only add if there are any contact details
        contact_p.add_run(contact_text)

    doc.add_paragraph("")  # small spacing

    # --------------------------------------------------------
    # PROFESSIONAL SUMMARY
    # --------------------------------------------------------

    add_section_header("Professional Summary")
    for point in resume_data["professional_summary"]:
        add_bullet_point(point)

    # --------------------------------------------------------
    # TECHNICAL SKILLS (ONLY TABLE USED)
    # --------------------------------------------------------

    add_section_header("Technical Skills")

    skills_table = doc.add_table(rows=0, cols=2)
    skills_table.autofit = False
    skills_table.columns[0].width = Inches(2)
    skills_table.columns[1].width = Inches(4.6)

    for category, values in resume_data["technical_skills"].items():
        row = skills_table.add_row()
        row.cells[0].paragraphs[0].add_run(category + ":").bold = True
        row.cells[1].paragraphs[0].add_run(", ".join(values))

    doc.add_paragraph("")

    # --------------------------------------------------------
    # EXPERIENCE (NO TABLES AT ALL)
    # --------------------------------------------------------

    add_section_header("Professional Experience")

    for exp in resume_data["experiences"]:

        # Line 1: Client Name   <tab>   Duration
        header_line = doc.add_paragraph()
        ts = header_line.paragraph_format.tab_stops
        ts.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)

        header_line.add_run(f"{exp['client_name']}\t{exp['duration']}").bold = True

        # Role
        role_p = doc.add_paragraph()
        role_p.add_run(f"Role: {exp['job_role']}").italic = True

        # Bullets
        for bullet in exp["responsibilities"]:
            add_bullet_point(bullet)

        # Environment
        environment = exp.get("environment")
        if environment:
            env_p = doc.add_paragraph()
            env_p.add_run("Environment: ").bold = True
            env_p.add_run(environment)

        doc.add_paragraph("")  # small spacing between experiences

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    add_section_header("Education")
    for edu in resume_data["education"]:
        p = doc.add_paragraph(edu)
        p.paragraph_format.space_after = Pt(0)

    # --------------------------------------------------------
    # SAVE FILE
    # --------------------------------------------------------

    filename = f"{resume_data['name'].replace(' ', '_')}_Resume.docx"
    doc.save(filename)

    return filename

# ------------------------------------------------------------
# 2. UPLOAD TO GCS
# ------------------------------------------------------------

def upload_to_gcs(bucket_name: str, file_path: str, destination_blob_name: str):
    storage_client = storage.Client.from_service_account_json("src/resume-477618-0c64e84c6bb0.json")
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(destination_blob_name)

    # Upload
    blob.upload_from_filename(file_path)

    # Public URL
    public_url = blob.public_url

    return public_url

# ------------------------------------------------------------
# 3. MASTER FUNCTION
# ------------------------------------------------------------
def generate_and_upload_resume(resume_data):
    local_file = create_resume(resume_data)
    gcs_url = upload_to_gcs(
        bucket_name="resume-ai-bucket",
        file_path=local_file,
        destination_blob_name=os.path.basename(local_file)
    )
    return {"gcs_url": gcs_url}